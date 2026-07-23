import asyncio
import base64
import io
import json
import mimetypes
import os
import time
from pprint import pprint
from textwrap import dedent

import docx
import httpx
import pypdf
from config.config import Settings
from dotenv import load_dotenv
from fastmcp import Client as FastMCPClient
from fastmcp.client.transports import StreamableHttpTransport
from google import genai
from google.genai import types
from schemas.request import QuestionRequest
from utils.prompt import system_prompt
from utils.tools_converter import (
    convert_mcp_tools_to_gemini,
    convert_mcp_tools_to_ollama,
)

from langfuse import get_client, observe, propagate_attributes
from supabase import Client as SupabaseClient
from supabase import create_client

# Models that use OpenAI-compatible API format
OLLAMA_MODELS = {"qwen3.5:27b"}
OPENAI_MODELS = {"gpt-4o", "gpt-4o-mini"}


class ChatRobot:

    def __init__(self):
        self.langfuse_client = get_client()
        self.settings = Settings()
        load_dotenv()
        self.gemini_client = genai.Client(api_key=self.settings.google_key)
        self.supabase: SupabaseClient = create_client(
            self.settings.supabase_url, self.settings.supabase_key
        )
        self.req = None

    # --- MODEL DETECTION ---

    def _is_openai_compatible(self, model_name: str) -> bool:
        """Check if the given model uses the OpenAI-compatible API (Ollama or OpenAI GPT)."""
        return model_name in OLLAMA_MODELS or model_name in OPENAI_MODELS

    def _model_supports_vision(self, model_name: str) -> bool:
        """Check if the given model supports multimodal vision input."""
        name_lower = model_name.lower()
        if "gemini" in name_lower:
            return True
        if "gpt-4" in name_lower:
            return True
        if "-vl" in name_lower or "llava" in name_lower:
            return True
        if "qwen3.5" in name_lower:
            return True
        return False

    def _get_openai_endpoint(self, model_name: str) -> tuple[str, dict]:
        """
        Returns (base_url, headers) for the given OpenAI-compatible model.
        - Ollama models → local Ollama server, no auth
        - OpenAI GPT models → api.openai.com, with Bearer token
        """
        if model_name in OPENAI_MODELS:
            return (
                "https://api.openai.com/v1/chat/completions",
                {
                    "Authorization": f"Bearer {self.settings.openai_api_key}",
                    "Content-Type": "application/json",
                },
            )
        else:
            # Ollama
            return (
                f"{self.settings.ollama_host}/v1/chat/completions",
                {"Content-Type": "application/json"},
            )

    # --- HISTORY MANAGEMENT ---

    @observe()
    def generate_session_title(
        self, session_id: str, user_prompt: str, bot_answer: str
    ):
        """
        Membuat judul sesi berdasarkan konteks percakapan pertama menggunakan Gemini.
        Hanya berjalan jika title belum ada.
        """
        try:
            current = (
                self.supabase.table("chat-sessions")
                .select("title")
                .eq("id", session_id)
                .execute()
            )

            if current.data and current.data[0].get("title"):
                print(f"   ℹ️ Session already has title: {current.data[0]['title']}")
                return

            print("   ✨ Generating smart title for this session...")

            title_prompt = (
                f"Berdasarkan percakapan berikut, buatkan judul sesi yang sangat singkat, "
                f"padat, dan deskriptif (maksimal 5 kata). Jangan gunakan tanda kutip.\n\n"
                f"User: {user_prompt}\n"
                f"Model: {bot_answer}"
            )

            resp = self.gemini_client.models.generate_content(
                model="gemini-2.5-flash", contents=title_prompt
            )

            new_title = resp.text.strip()

            self.supabase.table("chat-sessions").update({"title": new_title}).eq(
                "id", session_id
            ).execute()

            print(f"   🏷️ Title updated to: '{new_title}'")

        except Exception as e:
            print(f"   ⚠️ Failed to generate title: {e}")

    def create_history(self) -> str:
        """Membuat sesi baru di DB dan mengembalikan ID-nya."""
        res = self.supabase.table("chat-sessions").insert({}).execute()
        if res.data:
            new_id = res.data[0]["id"]
            print(f"🆕 Session Created: {new_id}")
            return new_id
        return None

    def save_message(
        self, session_id: str, content: types.Content, showed: bool = True
    ):
        """
        Menyimpan pesan ke DB.
        PENTING: Kita serialisasi objek Gemini ke JSON.
        Kita TIDAK menyimpan bytes file di sini.
        showed: True untuk user msg & final model text, False untuk function_call & function_response.
        """
        serialized_parts = []
        for part in content.parts:
            # Skip thinking/reasoning parts (internal chain-of-thought dari Gemini 2.5)
            if getattr(part, "thought", False):
                continue
            if part.text:
                serialized_parts.append({"text": part.text})
            elif part.function_call:
                serialized_parts.append(
                    {
                        "function_call": {
                            "name": part.function_call.name,
                            "args": (
                                dict(part.function_call.args)
                                if part.function_call.args
                                else {}
                            ),
                        }
                    }
                )
            elif part.function_response:
                serialized_parts.append(
                    {
                        "function_response": {
                            "name": part.function_response.name,
                            "response": part.function_response.response,
                        }
                    }
                )
            elif hasattr(part, "inline_data") and part.inline_data:
                b64_data = base64.b64encode(part.inline_data.data).decode("utf-8")
                serialized_parts.append(
                    {
                        "inline_data": {
                            "mime_type": part.inline_data.mime_type,
                            "data": b64_data,
                        }
                    }
                )

        # Safeguard: Jika serialized_parts kosong (karena safety block, empty/whitespace text, atau thought-only parts),
        # tambahkan fallback text agar history tidak kosong [] di database
        if not serialized_parts:
            fallback_text = (
                "[Empty Response]" if content.role == "model" else "[Empty Message]"
            )
            serialized_parts.append({"text": fallback_text})

        self.supabase.table("chat-messages").insert(
            {
                "session_id": session_id,
                "role": content.role,
                "content": serialized_parts,
                "showed": showed,
            }
        ).execute()

    def get_history(self, session_id: str) -> list:
        """
        Mengambil history dan melakukan RE-INJECTION file bytes jika diperlukan.
        """
        res = (
            self.supabase.table("chat-messages")
            .select("*")
            .eq("session_id", session_id)
            .order("created_at", desc=False)
            .execute()
        )

        gemini_messages = []

        for row in res.data:
            role = row["role"]
            db_content = row["content"]

            parts = []
            tool_response_payload = None

            for item in db_content:
                if "text" in item:
                    parts.append(types.Part.from_text(text=item["text"]))
                elif "function_call" in item:
                    fc = item["function_call"]
                    parts.append(
                        types.Part.from_function_call(name=fc["name"], args=fc["args"])
                    )
                elif "function_response" in item:
                    fr = item["function_response"]
                    tool_response_payload = fr["response"]
                    parts.append(
                        types.Part.from_function_response(
                            name=fr["name"], response=fr["response"]
                        )
                    )
                elif "inline_data" in item:
                    id_data = item["inline_data"]
                    file_bytes = base64.b64decode(id_data["data"])
                    parts.append(
                        types.Part.from_bytes(
                            data=file_bytes, mime_type=id_data["mime_type"]
                        )
                    )

            # Safeguard: Jika database content kosong [], buat fallback part agar tidak error di Gemini API
            if not parts:
                fallback_text = (
                    "[Empty Response]" if role == "model" else "[Empty Message]"
                )
                parts.append(types.Part.from_text(text=fallback_text))

            gemini_messages.append(types.Content(role=role, parts=parts))

            if role == "tool" and tool_response_payload:
                msg_type = tool_response_payload.get("type")
                status = tool_response_payload.get("status")

                if msg_type == "sop_query" and status == "success":
                    data = tool_response_payload.get("data", [])
                    for obj in data:
                        sop_url = obj.get("sop_url")
                        obj_name = obj.get("name", "unknown")
                        if sop_url:
                            print(
                                f"   📂 History Replay: Re-downloading SOP for '{obj_name}' from URL..."
                            )
                            file_injection_msg = self.download_file_from_url(
                                sop_url, obj_name
                            )
                            if file_injection_msg:
                                gemini_messages.append(file_injection_msg)

                elif msg_type == "image_capture" and status == "success":
                    data = tool_response_payload.get("data", {})
                    filepath = data.get("filepath")

                    print(
                        f"   📸 History Replay: Re-downloading captured image '{filepath}' for context..."
                    )

                    image_injection_msg = self.download_image(filepath)

                    if image_injection_msg:
                        gemini_messages.append(image_injection_msg)

        return gemini_messages

    def get_history_for_ollama(self, session_id: str, model_name: str = "") -> list:
        """
        Mengambil history dari DB dan convert ke format OpenAI messages untuk Ollama.
        File re-injection dilakukan sebagai text description (Ollama tidak support raw bytes).
        """
        res = (
            self.supabase.table("chat-messages")
            .select("*")
            .eq("session_id", session_id)
            .order("created_at", desc=False)
            .execute()
        )

        ollama_messages = []
        supports_vision = (
            self._model_supports_vision(model_name) if model_name else True
        )

        for row in res.data:
            role = row["role"]
            db_content = row["content"]

            for item in db_content:
                if "text" in item:
                    # Map Gemini roles to OpenAI roles
                    openai_role = "assistant" if role == "model" else role
                    ollama_messages.append(
                        {"role": openai_role, "content": item["text"]}
                    )
                elif "function_call" in item:
                    fc = item["function_call"]
                    # Reconstruct as assistant message with tool_calls
                    ollama_messages.append(
                        {
                            "role": "assistant",
                            "content": None,
                            "tool_calls": [
                                {
                                    "id": f"call_{fc['name']}",
                                    "type": "function",
                                    "function": {
                                        "name": fc["name"],
                                        "arguments": json.dumps(fc["args"]),
                                    },
                                }
                            ],
                        }
                    )
                elif "function_response" in item:
                    fr = item["function_response"]
                    ollama_messages.append(
                        {
                            "role": "tool",
                            "tool_call_id": f"call_{fr['name']}",
                            "content": json.dumps(fr["response"]),
                        }
                    )
                elif "inline_data" in item:
                    id_data = item["inline_data"]
                    mime = id_data["mime_type"]
                    b64_data = id_data["data"]
                    file_bytes = base64.b64decode(b64_data)

                    if mime.startswith("image/"):
                        if supports_vision:
                            if (
                                ollama_messages
                                and ollama_messages[-1]["role"] == "user"
                            ):
                                prev_content = ollama_messages[-1]["content"]
                                if isinstance(prev_content, str):
                                    ollama_messages[-1]["content"] = [
                                        {"type": "text", "text": prev_content},
                                        {
                                            "type": "image_url",
                                            "image_url": {
                                                "url": f"data:{mime};base64,{b64_data}"
                                            },
                                        },
                                    ]
                                elif isinstance(prev_content, list):
                                    ollama_messages[-1]["content"].append(
                                        {
                                            "type": "image_url",
                                            "image_url": {
                                                "url": f"data:{mime};base64,{b64_data}"
                                            },
                                        }
                                    )
                            else:
                                ollama_messages.append(
                                    {
                                        "role": "user",
                                        "content": [
                                            {
                                                "type": "image_url",
                                                "image_url": {
                                                    "url": f"data:{mime};base64,{b64_data}"
                                                },
                                            }
                                        ],
                                    }
                                )
                        else:
                            text_to_append = f"\n\n[System Injection] Image attached but vision is not supported by model '{model_name}'."
                            if (
                                ollama_messages
                                and ollama_messages[-1]["role"] == "user"
                            ):
                                if isinstance(ollama_messages[-1]["content"], str):
                                    ollama_messages[-1]["content"] += text_to_append
                                elif isinstance(ollama_messages[-1]["content"], list):
                                    ollama_messages[-1]["content"].append(
                                        {"type": "text", "text": text_to_append}
                                    )
                            else:
                                ollama_messages.append(
                                    {"role": "user", "content": text_to_append}
                                )
                    elif mime == "application/pdf":
                        try:
                            pdf_reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                            pages_text = []
                            for page in pdf_reader.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    pages_text.append(page_text)
                            extracted = (
                                "\n\n".join(pages_text)
                                if pages_text
                                else "(No extractable text found in PDF)"
                            )
                        except Exception as e:
                            extracted = f"(Error extracting PDF text: {e})"

                        text_to_append = f"\n\n[System Injection] Extracted PDF content:\n{extracted}"
                        if ollama_messages and ollama_messages[-1]["role"] == "user":
                            if isinstance(ollama_messages[-1]["content"], str):
                                ollama_messages[-1]["content"] += text_to_append
                            elif isinstance(ollama_messages[-1]["content"], list):
                                ollama_messages[-1]["content"].append(
                                    {"type": "text", "text": text_to_append}
                                )
                        else:
                            ollama_messages.append(
                                {"role": "user", "content": text_to_append}
                            )
                    elif (
                        mime
                        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ):
                        try:
                            doc_stream = io.BytesIO(file_bytes)
                            doc = docx.Document(doc_stream)
                            full_text = [para.text for para in doc.paragraphs]
                            extracted = "\n".join(full_text)
                        except Exception as e:
                            extracted = f"(Error extracting DOCX text: {e})"

                        text_to_append = f"\n\n[System Injection] Extracted DOCX content:\n{extracted}"
                        if ollama_messages and ollama_messages[-1]["role"] == "user":
                            if isinstance(ollama_messages[-1]["content"], str):
                                ollama_messages[-1]["content"] += text_to_append
                            elif isinstance(ollama_messages[-1]["content"], list):
                                ollama_messages[-1]["content"].append(
                                    {"type": "text", "text": text_to_append}
                                )
                        else:
                            ollama_messages.append(
                                {"role": "user", "content": text_to_append}
                            )
                    else:
                        text_to_append = f"\n\n[System Injection] Binary file attached ({mime}, {file_bytes} bytes). Cannot display content."
                        if ollama_messages and ollama_messages[-1]["role"] == "user":
                            if isinstance(ollama_messages[-1]["content"], str):
                                ollama_messages[-1]["content"] += text_to_append
                            elif isinstance(ollama_messages[-1]["content"], list):
                                ollama_messages[-1]["content"].append(
                                    {"type": "text", "text": text_to_append}
                                )
                        else:
                            ollama_messages.append(
                                {"role": "user", "content": text_to_append}
                            )

        return ollama_messages

    # --- TOOLS & HELPERS ---

    def download_file(self, file_name: str, folder_name: str) -> types.Content:
        """
        Mengunduh file. Jika PDF/Gambar -> kirim sebagai File Bytes.
        Jika DOCX -> ekstrak teksnya -> kirim sebagai Teks.
        """
        try:
            file_path = f"{folder_name}/{file_name}" if folder_name else file_name
            # 1. Download Bytes dari Supabase
            file_bytes = self.supabase.storage.from_(
                self.settings.bucket_name
            ).download(file_path)
            print(f"   📥 Downloaded file '{file_path}' ({len(file_bytes)} bytes)")
            if not file_bytes:
                return None

            # 2. Deteksi Mime Type
            mime_type, _ = mimetypes.guess_type(file_name)
            if mime_type is None:
                mime_type = "application/octet-stream"

            parts = []

            # 3. Cek apakah ini file DOCX (Word)
            if (
                mime_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                try:
                    # Gunakan io.BytesIO agar library docx bisa membaca raw bytes seolah-olah file fisik
                    doc_stream = io.BytesIO(file_bytes)
                    doc = docx.Document(doc_stream)

                    # Ekstrak semua paragraf menjadi satu string
                    full_text = [para.text for para in doc.paragraphs]
                    extracted_text = "\n".join(full_text)

                    # Masukkan sebagai Part TEXT
                    parts.append(
                        types.Part.from_text(
                            text=f"--- Content of {file_name} ---\n{extracted_text}"
                        )
                    )
                except Exception as e:
                    print(f"⚠️ Gagal parsing DOCX {file_name}: {e}")
                    # Fallback jika gagal parsing, kirim info error saja
                    parts.append(
                        types.Part.from_text(
                            text=f"Error reading docx content: {file_name}"
                        )
                    )

            # 4. Jika BUKAN docx (misal: PDF, Gambar, Video), kirim sebagai Bytes (Native)
            else:
                parts.append(
                    types.Part.from_bytes(data=file_bytes, mime_type=mime_type)
                )
                # Tambahkan label nama file (opsional, agar AI tahu nama filenya)
                parts.append(
                    types.Part.from_text(
                        text=f"[System Injection] File uploaded: {file_name}"
                    )
                )

            return types.Content(role="user", parts=parts)

        except Exception as e:
            print(f"❌ Error downloading file for history: {e}")
            return None

    def download_file_from_url(
        self, url: str, object_name: str = "unknown"
    ) -> types.Content:
        """
        Mengunduh file SOP dari URL lokal. Jika PDF/Gambar -> kirim sebagai File Bytes.
        Jika DOCX -> ekstrak teksnya -> kirim sebagai Teks.
        URL bersifat lokal dan tidak bisa diakses langsung oleh LLM.
        """
        try:
            import requests as req_lib

            print(f"   📥 Downloading SOP from URL: {url}")
            resp = req_lib.get(url, timeout=30)
            resp.raise_for_status()
            file_bytes = resp.content

            if not file_bytes:
                print(f"   ⚠️ File empty from URL: {url}")
                return None

            print(f"   📥 Downloaded SOP for '{object_name}' ({len(file_bytes)} bytes)")

            # Detect mime type from Content-Type header or URL
            content_type = resp.headers.get("Content-Type", "")
            if ";" in content_type:
                content_type = content_type.split(";")[0].strip()

            if not content_type or content_type == "application/octet-stream":
                # Fallback: guess from URL
                mime_type, _ = mimetypes.guess_type(url)
                if mime_type:
                    content_type = mime_type
                else:
                    content_type = "application/octet-stream"

            parts = []

            # DOCX -> extract text
            if (
                content_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                try:
                    doc_stream = io.BytesIO(file_bytes)
                    doc = docx.Document(doc_stream)
                    full_text = [para.text for para in doc.paragraphs]
                    extracted_text = "\n".join(full_text)
                    parts.append(
                        types.Part.from_text(
                            text=f"--- SOP Content for '{object_name}' ---\n{extracted_text}"
                        )
                    )
                except Exception as e:
                    print(f"   ⚠️ Gagal parsing DOCX dari URL: {e}")
                    parts.append(
                        types.Part.from_text(
                            text=f"Error reading docx content from URL for '{object_name}'"
                        )
                    )

            # PDF, Image, etc -> send as bytes
            else:
                parts.append(
                    types.Part.from_bytes(data=file_bytes, mime_type=content_type)
                )
                parts.append(
                    types.Part.from_text(
                        text=f"[System Injection] SOP file for object '{object_name}'"
                    )
                )

            return types.Content(role="user", parts=parts)

        except Exception as e:
            print(f"❌ Error downloading SOP from URL '{url}': {e}")
            return None

    def download_image(self, filepath: str) -> types.Content:
        """
        Mengunduh gambar hasil capture dari Supabase Storage dan mengembalikannya
        sebagai Content dengan image bytes untuk di-inject ke Gemini context.
        """
        try:
            # Download bytes dari Supabase Storage
            file_bytes = self.supabase.storage.from_(
                self.settings.bucket_name
            ).download(filepath)

            if not file_bytes:
                print(f"   ⚠️ Image file empty or not found: {filepath}")
                return None

            filename = filepath.split("/")[-1]

            parts = [
                types.Part.from_bytes(data=file_bytes, mime_type="image/jpeg"),
                types.Part.from_text(
                    text=f"[System Injection] Captured image from robot camera: {filename}"
                ),
            ]

            return types.Content(role="user", parts=parts)

        except Exception as e:
            print(f"❌ Error downloading captured image: {e}")
            return None

    def _parse_files(self, files: list[str]) -> list[tuple[bytes, str, str]]:
        """
        Parses a list of data URLs.
        Returns a list of tuples: (file_bytes, mime_type, filename)
        """
        parsed = []
        import re

        for file_data in files:
            if not file_data:
                continue
            # Check if it is a data URL: data:<mime>;base64,<data>
            match = re.match(r"^data:([^;]+);base64,(.*)$", file_data)
            if match:
                mime_type = match.group(1)
                base64_data = match.group(2)
                try:
                    file_bytes = base64.b64decode(base64_data)
                    ext = mimetypes.guess_extension(mime_type) or ".bin"
                    filename = f"uploaded_file{ext}"
                    parsed.append((file_bytes, mime_type, filename))
                except Exception as e:
                    print(f"⚠️ Failed to decode base64 file: {e}")
            else:
                print(f"⚠️ File data is not a base64 data URL: {file_data[:50]}...")
        return parsed

    def _build_gemini_user_message(
        self, user_prompt: str, parsed_files: list[tuple[bytes, str, str]]
    ) -> types.Content:
        parts = [types.Part.from_text(text=user_prompt)]
        for file_bytes, mime_type, filename in parsed_files:
            if mime_type.startswith("image/") or mime_type == "application/pdf":
                parts.append(
                    types.Part.from_bytes(data=file_bytes, mime_type=mime_type)
                )
            elif (
                mime_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                # DOCX extraction
                try:
                    doc_stream = io.BytesIO(file_bytes)
                    doc = docx.Document(doc_stream)
                    full_text = [para.text for para in doc.paragraphs]
                    extracted_text = "\n".join(full_text)
                    parts.append(
                        types.Part.from_text(
                            text=f"--- Content of {filename} ---\n{extracted_text}"
                        )
                    )
                except Exception as e:
                    print(f"⚠️ Failed to parse DOCX {filename}: {e}")
                    parts.append(
                        types.Part.from_text(
                            text=f"Error reading docx content: {filename}"
                        )
                    )
            else:
                parts.append(
                    types.Part.from_bytes(data=file_bytes, mime_type=mime_type)
                )
        return types.Content(role="user", parts=parts)

    def _build_openai_user_message(
        self,
        user_prompt: str,
        parsed_files: list[tuple[bytes, str, str]],
        model_name: str = "",
    ) -> dict:
        content_array = [{"type": "text", "text": user_prompt}]
        supports_vision = (
            self._model_supports_vision(model_name) if model_name else True
        )
        for file_bytes, mime_type, filename in parsed_files:
            b64_data = base64.b64encode(file_bytes).decode("utf-8")
            if mime_type.startswith("image/"):
                if supports_vision:
                    content_array.append(
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:{mime_type};base64,{b64_data}"},
                        }
                    )
                else:
                    content_array.append(
                        {
                            "type": "text",
                            "text": f"\n\n[System Injection] Image attached ({filename}) but vision is not supported by model '{model_name}'.",
                        }
                    )
            elif mime_type == "application/pdf":
                try:
                    pdf_reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                    pages_text = []
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pages_text.append(page_text)
                    extracted = (
                        "\n\n".join(pages_text)
                        if pages_text
                        else "(No extractable text found in PDF)"
                    )
                except Exception as e:
                    extracted = f"(Error extracting PDF text: {e})"
                content_array.append(
                    {
                        "type": "text",
                        "text": f"\n\n[System Injection] Extracted PDF content ({filename}):\n{extracted}",
                    }
                )
            elif (
                mime_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                # DOCX extraction for Ollama/OpenAI
                try:
                    doc_stream = io.BytesIO(file_bytes)
                    doc = docx.Document(doc_stream)
                    full_text = [para.text for para in doc.paragraphs]
                    extracted = "\n".join(full_text)
                except Exception as e:
                    extracted = f"(Error extracting DOCX text: {e})"
                content_array.append(
                    {
                        "type": "text",
                        "text": f"\n\n[System Injection] Extracted DOCX content ({filename}):\n{extracted}",
                    }
                )
            else:
                content_array.append(
                    {
                        "type": "text",
                        "text": f"\n\n[System Injection] Binary file attached ({mime_type}, {len(file_bytes)} bytes). Cannot display content.",
                    }
                )
        return {"role": "user", "content": content_array}

    # --- MAIN PROCESS ---

    @observe()
    async def _fetch_mcp_tools(self, client):
        await client.ping()
        tools_response = await client.list_tools()
        return tools_response

    @observe(as_type="generation")
    async def _call_gemini(self, messages, gemini_tools, index):
        model_name = (
            (self.req.model_name or "gemini-2.5-flash")
            if self.req
            else "gemini-2.5-flash"
        )
        # Log input and model before making the call
        self.langfuse_client.update_current_generation(
            input=f"[{len(messages)} messages context]",
            model=model_name,
            metadata={
                "tags": (
                    self.req.tags + [f"{index}"]
                    if self.req and self.req.tags
                    else [f"{index}"]
                )
            },
        )

        max_retries = 3
        sys_instruction = (
            self.req.system_instruction
            if (
                self.req
                and hasattr(self.req, "system_instruction")
                and self.req.system_instruction
            )
            else system_prompt
        )
        for attempt in range(max_retries):
            try:
                response = self.gemini_client.models.generate_content(
                    model=model_name,
                    contents=messages,
                    config=types.GenerateContentConfig(
                        tools=gemini_tools, system_instruction=sys_instruction
                    ),
                )

                # Trace token usage
                if hasattr(response, "usage_metadata") and response.usage_metadata:
                    self.langfuse_client.update_current_generation(
                        usage_details={
                            "input": getattr(
                                response.usage_metadata, "prompt_token_count", 0
                            ),
                            "output": getattr(
                                response.usage_metadata, "candidates_token_count", 0
                            ),
                        }
                    )

                return response
            except (httpx.ConnectError, httpx.TimeoutException) as e:
                if attempt < max_retries - 1:
                    wait_time = 2 ** (attempt + 1)  # 2s, 4s, 8s
                    print(
                        f"   ⚠️ Network error (attempt {attempt + 1}/{max_retries}): {e}"
                    )
                    print(f"   🔄 Retrying in {wait_time}s...")
                    await asyncio.sleep(wait_time)
                else:
                    print(f"   ❌ Network error persisted after {max_retries} attempts")
                    raise

    @observe(as_type="generation")
    async def _call_openai_compatible(
        self, messages: list[dict], tools: list[dict], index
    ) -> dict:
        """
        Call an OpenAI-compatible endpoint (Ollama or OpenAI GPT).
        Automatically routes to the correct URL and auth based on model name.

        Args:
            messages: List of OpenAI-format message dicts.
            tools: List of OpenAI-format tool definitions.

        Returns:
            The response JSON dict.
        """
        model_name = (
            (self.req.model_name or "qwen3.5:27b") if self.req else "qwen3.5:27b"
        )

        self.langfuse_client.update_current_generation(
            input=f"[{len(messages)} messages context]",
            model=model_name,
            metadata={
                "tags": (
                    self.req.tags + [f"{index}"]
                    if self.req and self.req.tags
                    else [f"{index}"]
                )
            },
        )

        url, headers = self._get_openai_endpoint(model_name)
        provider = "OpenAI" if model_name in OPENAI_MODELS else "Ollama"
        print(f"   🔗 Calling {provider} endpoint: {url} with model '{model_name}'")
        payload = {
            "model": model_name,
            "messages": messages,
            "temperature": 0.7,
        }
        # Only include tools if there are any
        if tools:
            payload["tools"] = tools

        max_retries = 3
        for attempt in range(max_retries):
            try:
                async with httpx.AsyncClient(timeout=180.0) as http_client:
                    resp = await http_client.post(url, json=payload, headers=headers)
                    resp.raise_for_status()
                    result = resp.json()

                # Trace token usage
                usage = result.get("usage", {})
                if usage:
                    self.langfuse_client.update_current_generation(
                        usage_details={
                            "input": usage.get("prompt_tokens", 0),
                            "output": usage.get("completion_tokens", 0),
                        }
                    )

                return result
            except (httpx.ConnectError, httpx.TimeoutException) as e:
                if attempt < max_retries - 1:
                    wait_time = 2 ** (attempt + 1)
                    print(
                        f"   ⚠️ {provider} network error (attempt {attempt + 1}/{max_retries}): {e}"
                    )
                    print(f"   🔄 Retrying in {wait_time}s...")
                    await asyncio.sleep(wait_time)
                else:
                    print(
                        f"   ❌ {provider} network error persisted after {max_retries} attempts"
                    )
                    raise

    @observe()
    async def _process_tool_call(self, client, tool_name, tool_args, session_id):
        runtime_file_injection = None
        is_async_running = False
        response_payload = {}
        metadata = {}
        metadata["session_id"] = str(session_id)
        metadata["trace_id"] = self.langfuse_client.get_current_trace_id()
        metadata["observation_id"] = self.langfuse_client.get_current_observation_id()
        metadata["model_name"] = (
            (self.req.model_name or "gemini-2.5-flash")
            if self.req
            else "gemini-2.5-flash"
        )

        print(f"🔧 Calling tool: {tool_name}({tool_args})")

        try:
            result = await client.call_tool(
                name=tool_name, arguments=tool_args, meta=metadata
            )
            raw_output = (
                result.content[0].text if hasattr(result, "content") else str(result)
            )
            print(f"🔧 Result tool: {tool_name}: {raw_output}")
            parsed_output = json.loads(raw_output)
            response_payload = parsed_output

            if isinstance(parsed_output, dict):
                msg_type = parsed_output.get("type")
                status = parsed_output.get("status")

                # Detect async tools that are still running
                if status == "running":
                    is_async_running = True
                    print(
                        f"   ⏳ Tool '{tool_name}' is async (status=running). Will break loop."
                    )

                if msg_type == "sop_query" and status == "success":
                    print(
                        "   📄 SOP detected. Downloading from URL for current context..."
                    )
                    data = parsed_output.get("data", [])
                    injection_parts = []
                    for obj in data:
                        sop_url = obj.get("sop_url")
                        obj_name = obj.get("name", "unknown")
                        if sop_url:
                            file_content = self.download_file_from_url(
                                sop_url, obj_name
                            )
                            if file_content:
                                injection_parts.extend(file_content.parts)
                    if injection_parts:
                        runtime_file_injection = types.Content(
                            role="user", parts=injection_parts
                        )

                # elif msg_type == "image_capture" and status == "success":
                #     print("   📸 Image captured. Downloading for current context...")
                #     data = parsed_output.get("data", {})
                #     filepath = data.get("filepath")
                #     runtime_file_injection = self.download_image(filepath)

        except Exception as e:
            print(f"   ❌ Error: {e}")
            response_payload = {"error": str(e)}

        return (
            tool_name,
            tool_args,
            response_payload,
            runtime_file_injection,
            is_async_running,
        )

    def _build_gemini_tool_msg(self, tool_name, response_payload):
        """Build a Gemini-format tool response Content object."""
        return types.Content(
            role="tool",
            parts=[
                types.Part.from_function_response(
                    name=tool_name, response=response_payload
                )
            ],
        )

    @observe()
    async def main(self, req: QuestionRequest):
        self.req = req
        return await self.process_chat()

    # --- GEMINI FLOW ---

    @observe()
    async def process_chat(self):
        model_name = (
            (self.req.model_name or "gemini-2.5-flash")
            if self.req
            else "gemini-2.5-flash"
        )

        # Route to OpenAI-compatible flow if the model is Ollama or GPT
        if self._is_openai_compatible(model_name):
            return await self.process_chat_openai_compatible()

        transport = StreamableHttpTransport(url=self.settings.mcp_url)
        client = FastMCPClient(transport)

        session_id = self.req.session_id
        messages = []

        if not session_id:
            session_id = self.create_history()
        else:
            print(f"📜 Loading history for session: {session_id}")
            messages = self.get_history(session_id)

        with propagate_attributes(
            session_id=str(session_id), tags=self.req.tags if self.req.tags else None
        ):
            parsed_files = self._parse_files(self.req.files) if self.req.files else []
            user_msg = self._build_gemini_user_message(
                self.req.user_prompt, parsed_files
            )
            messages.append(user_msg)
            self.save_message(session_id, user_msg)

            # Inject robot status as a function_call/function_response pair
            # so Gemini treats it as tool output (high attention) not user command
            # Must come AFTER user turn (Gemini requires function_call after user or function_response turn)
            if self.req.system_prompt:
                status_call_msg = types.Content(
                    role="model",
                    parts=[
                        types.Part.from_function_call(name="get_robot_status", args={})
                    ],
                )
                status_response_msg = types.Content(
                    role="tool",
                    parts=[
                        types.Part.from_function_response(
                            name="get_robot_status",
                            response={"status": self.req.system_prompt},
                        )
                    ],
                )
                messages.append(status_call_msg)
                messages.append(status_response_msg)
                self.save_message(session_id, status_call_msg, showed=False)
                self.save_message(session_id, status_response_msg, showed=False)

            print(f"User: {self.req.user_prompt}\n")

            async with client:
                mcp_tools_raw = await self._fetch_mcp_tools(client)
                gemini_tools = convert_mcp_tools_to_gemini(mcp_tools_raw)

                index = 0
                while True:
                    response = await self._call_gemini(messages, gemini_tools, index)
                    index += 1
                    candidate = response.candidates[0]

                    # Log finish_reason non-standard untuk debugging
                    if hasattr(candidate, "finish_reason") and candidate.finish_reason:
                        fr_str = str(candidate.finish_reason)
                        if "STOP" not in fr_str:
                            print(
                                f"   ⚠️ Gemini finish_reason: {candidate.finish_reason}"
                            )

                    # Guard: Gemini may return None parts (safety block, empty response)
                    if not candidate.content or not candidate.content.parts:
                        print(
                            "   ⚠️ Gemini returned empty response (no parts). Retrying..."
                        )
                        # Append a nudge so Gemini knows it needs to respond
                        messages.append(
                            types.Content(
                                role="user",
                                parts=[
                                    types.Part.from_text(
                                        text="[System] Your previous response was empty. Please try again."
                                    )
                                ],
                            )
                        )
                        continue

                    # Guard: Gemini 2.5 thinking model bisa mengirim response yang hanya berisi
                    # thinking parts (thought=True) tanpa visible text atau function calls.
                    # Jika ini terjadi, retry agar Gemini menghasilkan jawaban visible.
                    has_visible_content = any(
                        (p.text and not getattr(p, "thought", False)) or p.function_call
                        for p in candidate.content.parts
                    )
                    if not has_visible_content:
                        print(
                            f"   ⚠️ Gemini returned thinking-only response (no visible text/function_call). Retrying..."
                        )
                        messages.append(
                            types.Content(
                                role="user",
                                parts=[
                                    types.Part.from_text(
                                        text="[System] Your previous response contained only internal reasoning. Please provide your visible response."
                                    )
                                ],
                            )
                        )
                        continue

                    messages.append(candidate.content)

                    # Cek apakah response berisi function_call (tidak ditampilkan di UI)
                    has_function_call = any(
                        part.function_call for part in candidate.content.parts
                    )
                    self.save_message(
                        session_id, candidate.content, showed=not has_function_call
                    )

                    found_tool_call = False
                    hit_async_running = False

                    for part in candidate.content.parts:
                        if part.function_call:
                            found_tool_call = True
                            tool_name = part.function_call.name
                            tool_args = (
                                dict(part.function_call.args)
                                if part.function_call.args
                                else {}
                            )

                            (
                                tool_name,
                                tool_args,
                                response_payload,
                                runtime_file_injection,
                                is_async_running,
                            ) = await self._process_tool_call(
                                client=client,
                                tool_name=tool_name,
                                tool_args=tool_args,
                                session_id=session_id,
                            )

                            tool_msg = self._build_gemini_tool_msg(
                                tool_name, response_payload
                            )
                            messages.append(tool_msg)
                            self.save_message(session_id, tool_msg, showed=False)

                            if runtime_file_injection:
                                print(
                                    "   📎 Injecting file bytes to Gemini context (Runtime)..."
                                )
                                messages.append(runtime_file_injection)

                            if is_async_running:
                                hit_async_running = True
                                break  # Stop processing further tool calls in this turn

                    # If an async tool is running, break loop and wait for robot webhook callback
                    if hit_async_running:
                        # Extract any text the model said before the tool call (e.g. "Robot sedang menuju...")
                        model_text_parts = [
                            p.text for p in candidate.content.parts if p.text
                        ]
                        waiting_msg = (
                            model_text_parts[0]
                            if model_text_parts
                            else "Robot sedang menjalankan perintah..."
                        )
                        print(
                            f"   ⏳ Async tool running. Breaking loop. Waiting for robot webhook callback."
                        )

                        self.generate_session_title(
                            session_id=session_id,
                            user_prompt=self.req.user_prompt,
                            bot_answer=waiting_msg,
                        )

                        return {"session_id": session_id, "answer": waiting_msg}

                    if not found_tool_call:
                        if candidate.content.parts[0].text:
                            print(
                                f"\n✨ Final Response: {candidate.content.parts[0].text}"
                            )
                        break

            final_answer = messages[-1].parts[0].text if messages[-1].parts else ""

            self.generate_session_title(
                session_id=session_id,
                user_prompt=self.req.user_prompt,
                bot_answer=final_answer,
            )

            return {"session_id": session_id, "answer": final_answer}

    # --- OPENAI-COMPATIBLE FLOW (Ollama / GPT) ---

    @observe()
    async def process_chat_openai_compatible(self):
        """
        Agentic loop for OpenAI-compatible models (Ollama Qwen, OpenAI GPT, etc.).
        Handles tool injection and tool result in OpenAI message format,
        but persists history in Gemini format for DB consistency.
        """
        model_name = (
            (self.req.model_name or "qwen3.5:27b") if self.req else "qwen3.5:27b"
        )
        provider = "OpenAI" if model_name in OPENAI_MODELS else "Ollama"

        transport = StreamableHttpTransport(url=self.settings.mcp_url)
        client = FastMCPClient(transport)

        session_id = self.req.session_id
        openai_messages = []

        if not session_id:
            session_id = self.create_history()
        else:
            print(f"📜 Loading history for session ({provider}): {session_id}")
            openai_messages = self.get_history_for_ollama(session_id)

        with propagate_attributes(
            session_id=str(session_id), tags=self.req.tags if self.req.tags else None
        ):
            # Add system prompt as first message (OpenAI format)
            # Insert at position 0 so it's always first
            sys_instruction = (
                self.req.system_instruction
                if (
                    self.req
                    and hasattr(self.req, "system_instruction")
                    and self.req.system_instruction
                )
                else system_prompt
            )
            openai_messages.insert(0, {"role": "system", "content": sys_instruction})

            # Add user message with files
            parsed_files = self._parse_files(self.req.files) if self.req.files else []
            user_openai_msg = self._build_openai_user_message(
                self.req.user_prompt, parsed_files
            )
            openai_messages.append(user_openai_msg)

            # Save user message to DB in Gemini format
            user_gemini_msg = self._build_gemini_user_message(
                self.req.user_prompt, parsed_files
            )
            self.save_message(session_id, user_gemini_msg)

            # Inject robot status if provided (as assistant + tool message pair)
            if self.req.system_prompt:
                # Assistant message: simulates tool call
                status_call_id = "call_get_robot_status"
                status_assistant_msg = {
                    "role": "assistant",
                    "content": None,
                    "tool_calls": [
                        {
                            "id": status_call_id,
                            "type": "function",
                            "function": {"name": "get_robot_status", "arguments": "{}"},
                        }
                    ],
                }
                # Tool response message
                status_tool_msg = {
                    "role": "tool",
                    "tool_call_id": status_call_id,
                    "content": json.dumps({"status": self.req.system_prompt}),
                }
                openai_messages.append(status_assistant_msg)
                openai_messages.append(status_tool_msg)

                # Also save to DB in Gemini format
                gemini_call = types.Content(
                    role="model",
                    parts=[
                        types.Part.from_function_call(name="get_robot_status", args={})
                    ],
                )
                gemini_resp = types.Content(
                    role="tool",
                    parts=[
                        types.Part.from_function_response(
                            name="get_robot_status",
                            response={"status": self.req.system_prompt},
                        )
                    ],
                )
                self.save_message(session_id, gemini_call, showed=False)
                self.save_message(session_id, gemini_resp, showed=False)

            print(f"User: {self.req.user_prompt}\n")

            async with client:
                mcp_tools_raw = await self._fetch_mcp_tools(client)
                openai_tools = convert_mcp_tools_to_ollama(mcp_tools_raw)

                index = 0
                while True:
                    response = await self._call_openai_compatible(
                        openai_messages, openai_tools, index
                    )
                    index += 1
                    choice = response.get("choices", [{}])[0]
                    message = choice.get("message", {})
                    finish_reason = choice.get("finish_reason", "")

                    assistant_content = message.get("content", None)
                    tool_calls = message.get("tool_calls", None)

                    # Guard: empty response
                    if not assistant_content and not tool_calls:
                        print(f"   ⚠️ {provider} returned empty response. Retrying...")
                        openai_messages.append(
                            {
                                "role": "user",
                                "content": "[System] Your previous response was empty. Please try again.",
                            }
                        )
                        continue

                    # Append the full assistant message to conversation (including any tool_calls)
                    assistant_msg_for_history = {"role": "assistant"}
                    if assistant_content:
                        assistant_msg_for_history["content"] = assistant_content
                    else:
                        assistant_msg_for_history["content"] = None
                    if tool_calls:
                        assistant_msg_for_history["tool_calls"] = tool_calls
                    openai_messages.append(assistant_msg_for_history)

                    # Handle tool calls
                    if tool_calls:
                        # Save the assistant message with function_call(s) to DB (Gemini format)
                        gemini_parts = []
                        if assistant_content:
                            gemini_parts.append(
                                types.Part.from_text(text=assistant_content)
                            )
                        for tc in tool_calls:
                            func = tc.get("function", {})
                            tc_name = func.get("name", "")
                            tc_args_str = func.get("arguments", "{}")
                            try:
                                tc_args = (
                                    json.loads(tc_args_str)
                                    if isinstance(tc_args_str, str)
                                    else tc_args_str
                                )
                            except json.JSONDecodeError:
                                tc_args = {}
                            gemini_parts.append(
                                types.Part.from_function_call(
                                    name=tc_name, args=tc_args
                                )
                            )

                        gemini_assistant_content = types.Content(
                            role="model", parts=gemini_parts
                        )
                        self.save_message(
                            session_id, gemini_assistant_content, showed=False
                        )

                        hit_async_running = False

                        for tc in tool_calls:
                            func = tc.get("function", {})
                            tool_call_id = tc.get("id", "")
                            tool_name = func.get("name", "")
                            tool_args_str = func.get("arguments", "{}")

                            try:
                                tool_args = (
                                    json.loads(tool_args_str)
                                    if isinstance(tool_args_str, str)
                                    else tool_args_str
                                )
                            except json.JSONDecodeError:
                                tool_args = {}

                            (
                                tool_name,
                                tool_args,
                                response_payload,
                                runtime_file_injection,
                                is_async_running,
                            ) = await self._process_tool_call(
                                client=client,
                                tool_name=tool_name,
                                tool_args=tool_args,
                                session_id=session_id,
                            )

                            # Append tool result in OpenAI format
                            tool_result_msg = {
                                "role": "tool",
                                "tool_call_id": tool_call_id,
                                "content": json.dumps(response_payload),
                            }
                            openai_messages.append(tool_result_msg)

                            # Save to DB in Gemini format
                            gemini_tool_msg = self._build_gemini_tool_msg(
                                tool_name, response_payload
                            )
                            self.save_message(session_id, gemini_tool_msg, showed=False)

                            # File injection: handle both text and binary (inline_data) parts
                            if runtime_file_injection:
                                print(
                                    f"   📎 Injecting file content to {provider} context (Runtime)..."
                                )
                                content_array = []

                                for p in runtime_file_injection.parts:
                                    if p.text:
                                        content_array.append(
                                            {"type": "text", "text": p.text}
                                        )
                                    elif p.inline_data and p.inline_data.data:
                                        b64_data = base64.b64encode(
                                            p.inline_data.data
                                        ).decode("utf-8")
                                        mime = (
                                            p.inline_data.mime_type
                                            or "application/octet-stream"
                                        )

                                        if mime.startswith("image/"):
                                            # image_url supported by both GPT and Ollama
                                            content_array.append(
                                                {
                                                    "type": "image_url",
                                                    "image_url": {
                                                        "url": f"data:{mime};base64,{b64_data}"
                                                    },
                                                }
                                            )
                                        elif model_name in OPENAI_MODELS:
                                            # type: file only supported by OpenAI GPT
                                            content_array.append(
                                                {
                                                    "type": "file",
                                                    "file": {
                                                        "filename": "uploaded_document",
                                                        "file_data": f"data:{mime};base64,{b64_data}",
                                                    },
                                                }
                                            )
                                        else:
                                            # Ollama: no file support, extract text from PDF/DOCX
                                            if mime == "application/pdf":
                                                try:
                                                    pdf_reader = pypdf.PdfReader(
                                                        io.BytesIO(p.inline_data.data)
                                                    )
                                                    pages_text = []
                                                    for page in pdf_reader.pages:
                                                        page_text = page.extract_text()
                                                        if page_text:
                                                            pages_text.append(page_text)
                                                    extracted = (
                                                        "\n\n".join(pages_text)
                                                        if pages_text
                                                        else "(No extractable text found in PDF)"
                                                    )
                                                except Exception as e:
                                                    extracted = f"(Error extracting PDF text: {e})"
                                                content_array.append(
                                                    {
                                                        "type": "text",
                                                        "text": f"[System Injection] Extracted PDF content:\n{extracted}",
                                                    }
                                                )
                                            elif (
                                                mime
                                                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                            ):
                                                try:
                                                    doc_stream = io.BytesIO(
                                                        p.inline_data.data
                                                    )
                                                    doc = docx.Document(doc_stream)
                                                    full_text = [
                                                        para.text
                                                        for para in doc.paragraphs
                                                    ]
                                                    extracted = "\n".join(full_text)
                                                except Exception as e:
                                                    extracted = f"(Error extracting DOCX text: {e})"
                                                content_array.append(
                                                    {
                                                        "type": "text",
                                                        "text": f"[System Injection] Extracted DOCX content:\n{extracted}",
                                                    }
                                                )
                                            else:
                                                content_array.append(
                                                    {
                                                        "type": "text",
                                                        "text": f"[System Injection] Binary file attached ({mime}, {len(p.inline_data.data)} bytes). Cannot display content.",
                                                    }
                                                )

                                if content_array:
                                    openai_messages.append(
                                        {"role": "user", "content": content_array}
                                    )

                            if is_async_running:
                                hit_async_running = True
                                break

                        # Handle async tool
                        if hit_async_running:
                            waiting_msg = (
                                assistant_content
                                if assistant_content
                                else "Robot sedang menjalankan perintah..."
                            )
                            print(
                                f"   ⏳ Async tool running. Breaking loop ({provider})."
                            )

                            # self.generate_session_title(
                            #     session_id=session_id,
                            #     user_prompt=self.req.user_prompt,
                            #     bot_answer=waiting_msg,
                            # )

                            return {"session_id": session_id, "answer": waiting_msg}

                        # Continue loop to let model process tool results
                        continue

                    else:
                        # No tool calls — this is the final answer
                        final_text = assistant_content or ""
                        if final_text:
                            print(f"\n✨ Final Response ({provider}): {final_text}")

                        # Save final assistant text to DB in Gemini format
                        gemini_final = types.Content(
                            role="model", parts=[types.Part.from_text(text=final_text)]
                        )
                        self.save_message(session_id, gemini_final, showed=True)
                        break

            final_answer = assistant_content or ""

            # self.generate_session_title(
            #     session_id=session_id,
            #     user_prompt=self.req.user_prompt,
            #     bot_answer=final_answer,
            # )

            return {"session_id": session_id, "answer": final_answer}
